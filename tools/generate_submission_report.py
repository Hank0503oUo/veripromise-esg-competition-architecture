"""Generate Word and PDF technical report for organizer review."""
from __future__ import annotations

import math
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


PKG = Path(__file__).resolve().parents[1]
DOCX_OUT = PKG / "submission_v43_reproduction_report.docx"
PDF_OUT = PKG / "submission_v43_reproduction_report.pdf"
FIG_DIR = PKG / "figures"

FIGURES = {
    "algo": FIG_DIR / "full_algorithm_layer_map.png",
    "overview": FIG_DIR / "v43_pipeline_overview.png",
    "field": FIG_DIR / "v43_field_transplant.png",
    "gate": FIG_DIR / "v43_validation_gate.png",
}


TITLE = "AI CUP 2026 春季賽：ESG 永續承諾驗證競賽報告"


SECTIONS = [
    (
        "一、摘要",
        [
            "本隊最終提交採用欄位級整合、可歸因驗證與風險控管策略，最終檔案為 submission_v43_tlonly_v39d.csv，Private Leaderboard 分數為 0.6457201，排名 6/143。整體方法不是單一模型直接四欄輸出，而是先建立 v18/v18C/v25 的穩定欄位級 baseline，再只把最後驗證通過的 verification_timeline 改進移植到最終檔案。",
            "完整系統由五個演算法層組成：資料與文本正規化層、rule/domain feature 層、LLM logit/judge feature 層、LightGBM stacking/calibration 層，以及最後的欄位級 deployment gate。v43 的正式公式雖然很短：v43 = v25 base + v39d timeline-only override，但 v25 base 本身包含 v18/v18C prior stack、A/B/C 多引擎 logits、ticker empirical-Bayes prior、MTP/TEJ 特徵與 GLM-5.2 evidence_status judge；v39d 則是通過 gate 的 explicit target-year timeline rule。",
            "本報告依據專案 README、v18_latest_architecture.md、HARNESS/v43_repro_package/docs/v43_pipeline.md 與最後上傳紀錄整理。為利主辦方複現，最後 v43 組裝階段不需要 GPU、不需要呼叫 API，也不需要重新訓練大型模型；只需提供必要 deterministic artifacts，即可由 scripts/reproduce_v43.py 重建與驗證 SHA256。",
        ],
    ),
    (
        "二、系統方法與模型架構",
        [
            "系統以四欄任務為核心：promise_status、verification_timeline、evidence_status、evidence_quality。架構上先將官方 JSON、ESG 報告全文、公司/ticker/年度資訊轉成可被模型使用的 row-level feature table；再把規則訊號、公司年度訊號、LLM logits、ticker prior 組成 feature/logit bank；最後由 per-field LightGBM stack 進行多分類，並以 validation OOF 做 offset calibration 與 deployment gate。",
            "資料與規則層採 deterministic feature engineering。MTP/domain features 不是硬規則，而是將年份、百分比、數字密度、行動/政策/查證/含糊詞、GRI/ISO/SBT 等 cue 轉為 tabular features；TEJ company-year features 以 ticker/year 回接公司年度可觀測訊號；PDF/full text 只用於產生可檢索上下文與長文本推理輸入，不使用 test label。",
            "v18 prior stack 是主分類器骨架。它將 Engine A/B 的長文本 logits、MTP features、TEJ features、BERT/v6.3 feature bank 與 ticker prior 接成 per-field feature matrix。ticker prior 是 empirical-Bayes 平滑的 P(class|ticker) 加 log(count)：validation prior 只由 train_1000 計算以避免自洩漏，test prior 由 train+val 計算，符合部署條件。分類器使用 per-field LightGBM multiclass、class_weight balanced、5-fold GroupKFold by ticker、3 seeds ensemble，以及 coord-ascent class offsets。",
            "v18C 是多引擎擴充。它修復 K3 few-shot prompt 的 company key bug，使用 top-6 label-balanced few-shot retrieval 產生 Engine C logits，並與 Engine A/B 形成 A+B+C、3 engines × 3 seeds 的 ensemble。此層不是單純 API 投票，而是把 C 的 logits 和同一組 tabular/prior features 重新餵入 LightGBM stack，經 ablation 確認超過門檻後才部署。",
            "v25 與 v39d 是欄位級 deployment layer。v25 將 GLM-5.2 evidence_status judge 的 probability p 作為 evidence_status 的額外一維 feature，經同樣的 stack/calibration 後只移植 evidence_status 欄位。v39d 則不訓練模型，而是 explicit target-year anchor rule：只在 confidence >= 0.70、reason 為 explicit_target_year 或 explicit_target_year_range、base promise_status=Yes、且原 timeline 不是 N/A 時覆寫 timeline。最終 v43 只採用 v39d 的 verification_timeline，保留 v25 其餘三欄。",
        ],
    ),
    (
        "三、創新點",
        [
            "第一，採用欄位級可歸因提交策略。由於 leaderboard 回饋可逐欄觀察，本隊將每次改動限制在單一欄位或明確欄位組合，避免多個實驗同時混入造成歸因不清。v42/v43 即為 final-day 的 orthogonal probes：v42 只改 evidence_quality，v43 只改 verification_timeline。",
            "第二，timeline 改進使用 structurally safe explicit-year anchor。此規則只在文本具有明確目標年或目標年區間時啟動，並限制 promise_status 必須為 Yes，且原 timeline 不得為 N/A。這避免將無承諾或不適用樣本強行改為時間類別。",
            "第三，所有高風險 LLM/EQ 支線均需通過 gate 才能部署。即使大型語言模型可提供語義判斷，本隊仍以 validation、ablation、public probe 與格式檢查作為最後守門，未達門檻者不納入 private final。",
        ],
    ),
    (
        "四、資料處理",
        [
            "主要資料來源為主辦方提供之 train/validation/test JSON 與提交格式。前處理包含欄位標籤正規化、id 排序、N/A 字串保留、公司 ticker 對應、文字長度與數字/年份/關鍵詞特徵擷取，以及 rule-based MTP 特徵生成。",
            "另外使用隊伍整理之 ESG 報告文字、公司年度 ESG/TEJ 衍生訊號與報告切分後的 rule-based 特徵。這些資料用於特徵工程與模型輸入，不使用 test label。對於 final v43，必要 deterministic artifacts 為 submission_v25_evjudge.csv、v35_timeline_anchor_features_test.csv、v18C_test_probas.csv 與 v39d timeline source。",
            "輸出檔案皆經格式驗證：2000 rows、欄位順序固定為 id,promise_status,verification_timeline,evidence_status,evidence_quality、id 依數字排序、UTF-8 無 BOM、LF newline、N/A 不被轉成 NaN，並檢查各欄標籤是否屬於合法集合。",
        ],
    ),
    (
        "五、訓練方式",
        [
            "主要分類器使用 per-field LightGBM multiclass stack，而不是一個模型同時硬解四欄。每個欄位各自建立 feature matrix 與 label space，使用 GroupKFold by ticker 做 validation OOF，降低同公司語料在訓練/驗證之間互相污染的風險。LightGBM 使用 class_weight balanced，並沿用前期 Optuna 搜尋出的 v16 參數，最後以 3 個 random seeds 平均降低方差。",
            "校準層使用 OOF 機率做 class offset search。也就是先得到各欄各類別機率，再用 coord-ascent 搜尋類別偏置，使 validation weighted score / macro-F1 最佳。最後 decode 後才套用 cascade：promise_status=No 時 timeline/evidence/EQ 為 N/A；evidence_status=No 或 N/A 時 evidence_quality 為 N/A。不過在 field-level transplant 之後不再套 cascade，以免破壞 public 已驗證的欄位獨立性。",
            "LLM 輔助分支不直接取代最終分類器，而是轉成 feature 或候選分支。例如 Qwen/K3 LoRA 與 few-shot engines 輸出四欄 logits，作為 LightGBM 的輸入；GLM-5.2 evidence-status judge 輸出 probability p，只接到 evidence_status feature；REL9B EQ、frontier EQ judges、MTP-9B teacher 等則只做 ablation。未通過 validation gate 或 public probe 的分支不部署。",
            "最後 v39d/v43 階段不再訓練模型，而是 deterministic rule/column assembly。此設計使最終提交可在主辦方環境中不依賴 GPU 或 API 重跑。",
        ],
    ),
    (
        "六、結果分析與討論",
        [
            "v25 base public score 為 0.6060570，當時已為穩定基準。最後一天針對 timeline 與 EQ 進行單欄 probe。EQ 類支線雖能在部分 validation 切分上改善，但 public 或 gate 顯示風險偏高，因此未進入 private final。timeline v39d 只改 57 筆並保持 cascade 結構，風險較低且效果明確。",
            "值得注意的是，本地 validation / gate 的估計分數比 public leaderboard 更接近最終 private 結果。HARNESS 與 outputs 紀錄顯示：v18 prior stack validation WS 為 0.6412508，v18C A+B+C validation WS 為 0.6506880，早期 v5.3 baseline 也記錄 weighted score 0.6460；最終 private score 為 0.6457201，落在本地估計帶中間。相較之下，最終提交在 public rank 僅為 59，但 private rank 提升至 6/143，顯示 public split 可能低估了本方法，且本隊的 grouped validation、單欄 ablation 與 deployment gate 更接近 private distribution。",
            "最終 v43 欄位分布為：promise_status Yes=1617、No=383；verification_timeline already=580、between_2_and_5_years=612、more_than_5_years=330、N/A=383、within_2_years=95；evidence_status Yes=1309、No=286、N/A=405；evidence_quality Clear=1154、N/A=700、Not Clear=146。v43 相對 v25 只改 timeline 57 筆。",
            "Private Leaderboard 最終結果為 0.6457201，排名 6/143。此結果顯示欄位級風險控管與 structurally safe timeline override 是有效的最後提交策略。",
        ],
    ),
    (
        "七、程式碼",
        [
            "複現程式碼位於 esg_competition/HARNESS/v43_repro_package。主要檔案包含 scripts/build_v39d_timeline.py、scripts/reproduce_v43.py、colab/reproduce_v43_colab.ipynb、README.md、MANIFEST.md 與 docs/v43_pipeline.md。",
            "執行方式：先安裝 requirements.txt，再執行 python scripts/build_v39d_timeline.py --root <ESG競賽根目錄>，接著執行 python scripts/reproduce_v43.py --root <ESG競賽根目錄>。後者會輸出 v43_repro_report.json，並檢查 SHA256 是否等於最終提交檔。",
            "公開 GitHub 架構文件位於 https://github.com/Hank0503oUo/veripromise-esg-competition-architecture ，其中 README.md、v18架構圖.md、v43最終架構.md 說明公開可揭露的演算法架構與驗證紀律。公開 repo 不放主辦方資料、submission CSV、LLM API 輸出、模型權重或大型中間 artifacts；必要 deterministic artifacts 可依主辦方要求以私有雲端資料夾提供。",
        ],
    ),
    (
        "八、使用的外部資源與參考文獻",
        [
            "外部資料與模型使用揭露：GLM-5.2（zai/glm-5.2）用於 evidence_status judge probability feature；Qwen3.5 9B / K3 LoRA 或 few-shot engine 用於上游 v18C 多引擎 logits；OpenAI Codex、Claude Opus、DeepSeek V4 Pro 等生成式 AI 助手用於程式撰寫協助、實驗規劃、錯誤排查與文件整理。GPT 產生之 Misleading hard-negative、MTP-9B teacher、REL9B EQ 與其他 frontier judge 屬於實驗支線，未直接形成 v43 的 evidence_quality 覆蓋。",
            "貢獻比例誠實聲明：以下為審查用估計，非數學精確分攤。隊伍人工設計、驗證與提交決策約 60%；傳統 ML/特徵工程/欄位級 ensemble 程式約 25%；生成式 AI 產生 judge 特徵或輔助標註約 10%；生成式 AI 對程式碼撰寫、文件與除錯的輔助約 5%。最終提交由隊伍依 validation、public probe 與風險控管決定。",
            "參考文獻：Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T.-Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems. Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830. Harris, C. R., et al. (2020). Array programming with NumPy. Nature, 585, 357-362. McKinney, W. (2010). Data structures for statistical computing in Python. Proceedings of the 9th Python in Science Conference.",
        ],
    ),
]


TABLES = {
    "team": [
        ["隊伍名稱", "TEAM_hank050389（請依 AI CUP 報名系統確認）"],
        ["隊員", "hank050389（其餘姓名、學校/系所、電話、E-mail 請依報名資料補齊）"],
        ["指導教授 / 業師", "無正式指導教授或業師；如有，請補姓名、學校/公司、科系/職稱、信箱"],
        ["Private leaderboard", "0.6457201 / Rank 6/143"],
    ],
    "v43_vs_v25": [
        ["Field", "Changed cells vs v25"],
        ["promise_status", "0"],
        ["verification_timeline", "57"],
        ["evidence_status", "0"],
        ["evidence_quality", "0"],
    ],
    "checksums": [
        ["Artifact", "SHA256"],
        ["submission_v25_evjudge.csv", "2da16c99e9b740ff7b9e638791f659fed246a4483e2fac4edcb913c6fca641a0"],
        ["submission_v39d_timeline_explicit_year_tau070_structsafe.csv", "17fd3822058bf413b4b0475881af0268b88a1a3d290f29376ef1b06b124c896a"],
        ["submission_v43_tlonly_v39d.csv", "17fd3822058bf413b4b0475881af0268b88a1a3d290f29376ef1b06b124c896a"],
    ],
    "external_ai": [
        ["Tool / Model", "Use", "Final contribution"],
        ["zai/glm-5.2", "evidence_status judge probability feature", "Included in v25 base"],
        ["Qwen3.5 9B / K3 LoRA/few-shot", "upstream multi-engine logits/features", "Included only through validated stack"],
        ["OpenAI Codex / Claude / DeepSeek", "coding, debugging, experiment planning, documentation", "Auxiliary; team validated final"],
        ["GPT Misleading 60 / REL9B EQ / MTP-9B / frontier judges", "candidate branches and ablations", "Rejected or not directly used in v43"],
    ],
    "github_architecture_map": [
        ["GitHub architecture section", "Meaning in this report", "Algorithmic structure"],
        ["README 架構總覽 / v18架構圖.md §1", "Overall system pipeline", "Data sources -> feature layer -> LLM/judge layer -> per-field LightGBM stack -> field-best output"],
        ["v18架構圖.md §3 檔案與資料層", "Data and feature normalization", "Official sample normalization, PDF/page/chunk context, TEJ company-year join, ticker/year alignment"],
        ["README 模型層 §1-§4", "Feature layer", "Transformer/logit features, regex/MTP domain cues, context retrieval features, empirical-Bayes ticker prior"],
        ["README 模型層 §5 / v18架構圖.md §5", "LLM feature layer", "Qwen/K3 LoRA logits, few-shot Engine C, GLM-5.2 narrow judge probability features"],
        ["README 集成邏輯 / v18架構圖.md §4", "Fusion and final selection", "Multi-engine averaging, LightGBM multiclass stack, OOF offset calibration, field-level transplant"],
        ["README 驗證紀律", "Deployment gate", "Grouped validation, single-field ablation, public probe, rejected branches kept out of final"],
    ],
    "algorithm_layers": [
        ["Layer", "Algorithm / method", "Where it is used"],
        ["Data normalization", "Deterministic JSON/CSV normalization, id sorting, N/A preservation, ticker/year/company alignment", "All train/val/test rows before feature construction"],
        ["Domain features", "Rule/MTP feature extraction: year, percent, quantity, assurance/ISO/GRI/SBT cues, vague/action words, numeric/table-like scores", "Tabular feature bank for all four fields"],
        ["Company-year features", "TEJ/ESG company-year join by ticker and year, with fallback to nearby year when needed", "Adds structured company context to LightGBM"],
        ["Ticker prior", "Empirical-Bayes smoothed P(label|ticker) + log sample count, alpha=2.0", "v18/v18C/v25 stack; validation computed train-only, test computed train+val"],
        ["LLM logits", "Qwen/K3 LoRA option logits and Engine A/B/C long-text scores", "Semantic probability features consumed by LightGBM, not direct final labels"],
        ["Few-shot Engine C", "Company prompt repair + top-6 label-balanced retrieval for examples", "Adds diversity to v18C A+B+C ensemble"],
        ["Stacking model", "Per-field LightGBM multiclass with class_weight balanced, 5-fold GroupKFold by ticker, 3-seed averaging, v16 Optuna params", "Main supervised fusion model"],
        ["Calibration", "OOF probability decode plus coordinate-ascent class offsets", "Optimizes validation macro-F1 / weighted score before deployment"],
        ["Postprocess", "Cascade during normal decode; no cascade after field-level transplant", "Keeps legal outputs while preserving independently validated field gains"],
        ["v25 judge feature", "GLM-5.2 evidence_status probability p appended as one feature", "Improves evidence_status only"],
        ["v39d timeline rule", "Explicit target-year anchor, confidence >= 0.70, structural guards", "Improves verification_timeline only"],
    ],
    "field_algorithm": [
        ["Field", "Main algorithmic source", "Reason"],
        ["promise_status", "v18C A/B/C long-text logits + LightGBM stack + ticker prior", "Engine C improved promise signal and passed ablation gate"],
        ["verification_timeline", "v39d explicit-year anchor over v25 base", "Timeline has concrete year/range cues; 57 structurally safe replacements"],
        ["evidence_status", "v25 LightGBM stack with GLM-5.2 judge probability feature", "Narrow LLM judge improved this field without touching others"],
        ["evidence_quality", "v18/v18C conservative stack", "EQ relation probes were unstable; incumbent stack was safest"],
    ],
    "local_private_alignment": [
        ["Evidence", "Score / rank", "Interpretation"],
        ["v5.3 local baseline", "weighted score 0.6460", "Early local validation already estimated the task around the final private range"],
        ["v18 prior stack", "val WS 0.6412508", "Stable grouped validation lower-bound candidate"],
        ["v18C A+B+C ensemble", "val WS 0.6506880", "Best local validation band after Engine C"],
        ["v25 public known score", "public WS 0.6060570", "Public leaderboard appeared lower than local/private estimate"],
        ["v43 final", "private 0.6457201, rank 6/143; public rank 59", "Private result landed inside the local validation band and jumped 53 ranks"],
    ],
    "script_map": [
        ["Component", "Main files", "Role in final explanation"],
        ["v18 prior stack", "build_v18_prior.py, integrate_v16_real.py, predict_v6_3.py", "LightGBM stack with MTP/TEJ/BERT/logit/ticker-prior features"],
        ["v18C engine-C ensemble", "build_v18C_engine3.py, K3 few-shot prompt artifacts", "Adds repaired company key and balanced few-shot Engine C logits"],
        ["v25 ev_status deployment", "build_v25_evjudge_deploy.py", "Transplants validated evidence_status improvement into field-best base"],
        ["v39d timeline probe", "build_v39_lastday_candidates.py, scripts/build_v39d_timeline.py", "Builds explicit-year timeline override with structural safety"],
        ["v43 final reproduction", "scripts/reproduce_v43.py, colab/reproduce_v43_colab.ipynb", "Deterministically assembles final CSV and verifies checksum"],
    ],
    "rejected": [
        ["Branch", "Reason not included in v43"],
        ["v19 retrieval / RAG features", "Redundant with ticker prior; validation gate decreased WS"],
        ["v20 local 9B judge / MTP-9B teacher", "Did not beat v18C/v25 after leakage-safe validation"],
        ["v21/v22 EQ binary/rubric", "EQ gains below gate or noisy; not worth private risk"],
        ["REL9B / v30 EQ blend / frontier EQ judges", "Public/gate showed EQ risk; v43 keeps incumbent EQ"],
        ["Misleading similarity probe", "Known positives required too large top-k; expected value negative"],
    ],
}


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/msjhbd.ttc" if bold else "C:/Windows/Fonts/msjh.ttc",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for item in candidates:
        path = Path(item)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def _wrap_label(text: str, width: int = 22) -> str:
    lines: list[str] = []
    for raw_line in text.split("\n"):
        if len(raw_line) <= width:
            lines.append(raw_line)
        else:
            lines.extend(textwrap.wrap(raw_line, width=width, break_long_words=False))
    return "\n".join(lines)


def _center_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.ImageFont, fill: str = "#111827") -> None:
    font_size = int(getattr(font, "size", 22))
    box_width = max(80, box[2] - box[0] - 28)
    wrap_width = max(12, int(box_width / max(8, font_size * 0.55)))
    wrapped = _wrap_label(text, width=wrap_width)
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=4, align="center")
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    x = box[0] + (box[2] - box[0] - tw) / 2
    y = box[1] + (box[3] - box[1] - th) / 2
    draw.multiline_text((x, y), wrapped, font=font, fill=fill, spacing=4, align="center")


def _box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str,
    outline: str = "#334155",
    font: ImageFont.ImageFont | None = None,
    radius: int = 14,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    _center_text(draw, box, text, font or _font(26, bold=True))


def _arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#0f766e", width: int = 5) -> None:
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 16
    points = [
        end,
        (end[0] - size * math.cos(angle - math.pi / 7), end[1] - size * math.sin(angle - math.pi / 7)),
        (end[0] - size * math.cos(angle + math.pi / 7), end[1] - size * math.sin(angle + math.pi / 7)),
    ]
    draw.polygon(points, fill=color)


def make_figures() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    title_font = _font(38, bold=True)
    head_font = _font(25, bold=True)
    body_font = _font(22)
    small_font = _font(18)

    # Figure 0: full algorithm layer map, aligned with the public GitHub architecture.
    img = Image.new("RGB", (1900, 1260), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.text((70, 42), "Figure 0. Full algorithm layer map", font=title_font, fill="#0f172a")
    lane_font = _font(24, bold=True)
    layer_font = _font(22, bold=True)
    desc_font = _font(18)
    y = 145
    lanes = [
        (
            "Data layer",
            "#dbeafe",
            [
                ("Official train/val/test", "label schema, id/order, N/A preservation"),
                ("ESG reports + TEJ", "PDF/page/chunk context, ticker-year company signals"),
            ],
        ),
        (
            "Feature layer",
            "#dcfce7",
            [
                ("Rule/MTP features", "years, numbers, assurance, ISO/GRI/SBT, vague/action cues"),
                ("Ticker prior", "smoothed P(label|ticker) + log(count), leakage-safe split"),
                ("Context features", "page/chunk retrieval only when field-specific"),
            ],
        ),
        (
            "LLM/logit layer",
            "#ede9fe",
            [
                ("Engine A/B/C logits", "Qwen/K3 LoRA option probabilities as features"),
                ("Engine C few-shot", "company-key repair + top-6 label-balanced examples"),
                ("GLM-5.2 judge", "narrow evidence_status probability feature"),
            ],
        ),
        (
            "Stack + calibration",
            "#fef3c7",
            [
                ("Per-field LightGBM", "multiclass, class_weight balanced, Optuna params"),
                ("GroupKFold OOF", "ticker-group validation + 3-seed averaging"),
                ("Offset calibration", "coordinate-ascent class bias on OOF probabilities"),
            ],
        ),
        (
            "Deployment layer",
            "#ffe4e6",
            [
                ("Field-best transplant", "replace only independently validated fields"),
                ("v39d timeline rule", "explicit-year anchor, confidence >= 0.70, structural guards"),
                ("Gate discipline", "single-field ablation, public probe, rejected branches excluded"),
            ],
        ),
    ]
    for lane_idx, (lane_name, color, items) in enumerate(lanes):
        lane_y = y + lane_idx * 210
        _box(draw, (70, lane_y, 310, lane_y + 155), lane_name, color, font=lane_font)
        last_right = 310
        for item_idx, (name, desc) in enumerate(items):
            x0 = 390 + item_idx * 480
            box = (x0, lane_y, x0 + 390, lane_y + 155)
            _box(draw, box, f"{name}\n{desc}", "#ffffff", font=desc_font)
            if item_idx == 0:
                _arrow(draw, (310, lane_y + 78), (x0, lane_y + 78), "#64748b", width=4)
            else:
                _arrow(draw, (last_right, lane_y + 78), (x0, lane_y + 78), "#94a3b8", width=3)
            last_right = x0 + 390
        if lane_idx < len(lanes) - 1:
            _arrow(draw, (1650, lane_y + 155), (1650, lane_y + 210), "#0f766e", width=4)
    _center_text(
        draw,
        (90, 1170, 1810, 1230),
        "This diagram follows the public GitHub architecture: data sources -> feature layer -> LLM/judge layer -> field-level LightGBM stack -> deployment gate.",
        small_font,
        "#334155",
    )
    img.save(FIGURES["algo"], quality=95)

    # Figure 1: overall architecture.
    img = Image.new("RGB", (1900, 1080), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.text((70, 42), "Figure 1. v43 overall pipeline architecture", font=title_font, fill="#0f172a")
    boxes = {
        "raw": (70, 160, 360, 330),
        "feat": (440, 160, 730, 330),
        "stack": (810, 160, 1100, 330),
        "base": (1180, 160, 1470, 330),
        "final": (1550, 160, 1840, 330),
        "v39d": (1180, 520, 1470, 700),
        "reject": (810, 760, 1470, 940),
    }
    _box(draw, boxes["raw"], "Official JSON\nESG reports\nTEJ company-year", "#dbeafe", font=head_font)
    _box(draw, boxes["feat"], "Feature/logit bank\nMTP, ticker prior\nEngine A/B/C", "#dcfce7", font=head_font)
    _box(draw, boxes["stack"], "v18 / v18C\nLightGBM stack\nOOF calibration", "#fef3c7", font=head_font)
    _box(draw, boxes["base"], "v25 base\nGLM-5.2 ev_status\nfieldbest", "#ede9fe", font=head_font)
    _box(draw, boxes["final"], "v43 final\nPrivate 0.6457201\nRank 6/143", "#ffe4e6", font=head_font)
    _box(draw, boxes["v39d"], "v39d timeline-only\nexplicit target-year\n57 safe overrides", "#ccfbf1", font=head_font)
    _box(draw, boxes["reject"], "Rejected branches: EQ probes, RAG features, MTP-9B,\nREL9B blend, frontier judges, Misleading similarity\nAll blocked by validation gate or public probe", "#e2e8f0", font=body_font)
    _arrow(draw, (360, 245), (440, 245))
    _arrow(draw, (730, 245), (810, 245))
    _arrow(draw, (1100, 245), (1180, 245))
    _arrow(draw, (1470, 245), (1550, 245))
    _arrow(draw, (1325, 520), (1325, 330))
    _arrow(draw, (1470, 610), (1620, 330))
    _arrow(draw, (1140, 760), (1325, 700), color="#64748b")
    _center_text(draw, (80, 990, 1820, 1045), "Final assembly is deterministic: v43 = v25 base + verification_timeline from v39d. No API/GPU is needed in the reproduction step.", small_font, "#334155")
    img.save(FIGURES["overview"], quality=95)

    # Figure 2: field-level source map.
    img = Image.new("RGB", (1900, 1080), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((70, 42), "Figure 2. Field-level transplant used by v43", font=title_font, fill="#0f172a")
    left_x, mid_x, right_x = 80, 660, 1240
    y0, row_h, gap = 160, 150, 34
    rows = [
        ("promise_status", "v18C engine-C ensemble", "unchanged in v43", "weight 0.20"),
        ("verification_timeline", "v39d explicit-year anchor", "57 cells replaced", "weight 0.15"),
        ("evidence_status", "v25 GLM-5.2 evjudge stack", "unchanged in v43", "weight 0.30"),
        ("evidence_quality", "v18/v18C incumbent EQ", "unchanged in v43", "weight 0.35"),
    ]
    for i, (field, source, action, weight) in enumerate(rows):
        y = y0 + i * (row_h + gap)
        _box(draw, (left_x, y, left_x + 420, y + row_h), field + "\n" + weight, "#eff6ff", font=head_font)
        _box(draw, (mid_x, y, mid_x + 430, y + row_h), source, "#ecfdf5", font=head_font)
        _box(draw, (right_x, y, right_x + 440, y + row_h), action, "#fff7ed", font=head_font)
        _arrow(draw, (left_x + 420, y + row_h // 2), (mid_x, y + row_h // 2), "#2563eb")
        _arrow(draw, (mid_x + 430, y + row_h // 2), (right_x, y + row_h // 2), "#2563eb")
    _box(draw, (520, 910, 1380, 1010), "No cascade after field transplant. Field scores transfer independently, so accepted improvements can be merged one column at a time.", "#f1f5f9", font=body_font)
    img.save(FIGURES["field"], quality=95)

    # Figure 3: validation gate.
    img = Image.new("RGB", (1900, 1080), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.text((70, 42), "Figure 3. Validation and deployment gate", font=title_font, fill="#0f172a")
    nodes = {
        "cand": (90, 180, 410, 340),
        "single": (560, 180, 880, 340),
        "gate": (1030, 180, 1350, 340),
        "deploy": (1500, 180, 1820, 340),
        "reject": (1030, 610, 1820, 830),
        "safe": (560, 610, 880, 830),
        "docs": (90, 610, 410, 830),
    }
    _box(draw, nodes["cand"], "New branch\nRAG / EQ / LLM\nTimeline", "#dbeafe", font=head_font)
    _box(draw, nodes["single"], "Single-field\nablation\nno mixed changes", "#dcfce7", font=head_font)
    _box(draw, nodes["gate"], "Gate check\nvalidation + public probe\nformat + cascade", "#fef3c7", font=head_font)
    _box(draw, nodes["deploy"], "Deploy only if\nfield gain is clear\nand risk is low", "#bbf7d0", font=head_font)
    _box(draw, nodes["reject"], "Do not merge if gate fails:\nv19 RAG, v20 9B, v21/v22 EQ,\nREL9B/v30, frontier EQ judges,\nMisleading similarity probe", "#fecaca", font=body_font)
    _box(draw, nodes["safe"], "Accepted final branch:\nv39d timeline-only\nstructural guards", "#ccfbf1", font=head_font)
    _box(draw, nodes["docs"], "Document artifacts\nREADME, MANIFEST,\nchecksums, Colab", "#e0e7ff", font=head_font)
    _arrow(draw, (410, 260), (560, 260))
    _arrow(draw, (880, 260), (1030, 260))
    _arrow(draw, (1350, 260), (1500, 260))
    _arrow(draw, (1190, 340), (1190, 610), "#dc2626")
    _arrow(draw, (720, 340), (720, 610), "#0f766e")
    _arrow(draw, (560, 720), (410, 720), "#4f46e5")
    _center_text(draw, (90, 930, 1820, 1010), "This gate is why v43 stayed small: only the timeline branch passed the risk/benefit threshold for the private-counting file.", body_font, "#334155")
    img.save(FIGURES["gate"], quality=95)


def add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = rows[0][i]
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def add_docx_figure(doc: Document, path: Path, caption: str, width_inches: float = 6.25) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9)
    doc.add_picture(str(path), width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.18)

    title = doc.add_heading(TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("隊伍與成績資訊", level=1)
    add_docx_table(doc, TABLES["team"])

    for heading, paragraphs in SECTIONS:
        doc.add_heading(heading, level=1)
        for paragraph in paragraphs:
            p = doc.add_paragraph(paragraph)
            p.paragraph_format.space_after = Pt(6)
        if heading.startswith("二、"):
            doc.add_heading("GitHub 架構文章對應", level=2)
            add_docx_table(doc, TABLES["github_architecture_map"])
            doc.add_heading("完整演算法層架構圖", level=2)
            add_docx_figure(doc, FIGURES["algo"], "圖 0：依 GitHub 公開架構整理的完整演算法層，從資料層一路到 deployment gate。")
            doc.add_heading("演算法層對照表", level=2)
            add_docx_table(doc, TABLES["algorithm_layers"])
            doc.add_heading("最終欄位來源與演算法理由", level=2)
            add_docx_table(doc, TABLES["field_algorithm"])
            doc.add_heading("最終組裝流程圖", level=2)
            add_docx_figure(doc, FIGURES["overview"], "圖 1：完整模型收斂到 v25 base，再由 v39d timeline-only 分支形成最終 v43。")
            add_docx_figure(doc, FIGURES["field"], "圖 2：field-best assembly：四欄各自選擇通過 gate 的最佳演算法來源。")
            doc.add_heading("v43 vs v25 欄位異動", level=2)
            add_docx_table(doc, TABLES["v43_vs_v25"])
        if heading.startswith("三、"):
            doc.add_heading("驗證與部署 gate", level=2)
            add_docx_figure(doc, FIGURES["gate"], "圖 3：所有新分支先做單欄 ablation 與 gate，未通過者不併入 private final。")
            doc.add_heading("未納入 v43 的支線", level=2)
            add_docx_table(doc, TABLES["rejected"])
        if heading.startswith("六、"):
            doc.add_heading("本地驗證與 private 結果對齊", level=2)
            add_docx_table(doc, TABLES["local_private_alignment"])
        if heading.startswith("七、"):
            doc.add_heading("主要腳本與角色", level=2)
            add_docx_table(doc, TABLES["script_map"])
            doc.add_heading("必要 artifact checksum", level=2)
            add_docx_table(doc, TABLES["checksums"])
        if heading.startswith("八、"):
            doc.add_heading("生成式 AI / 外部模型使用摘要表", level=2)
            add_docx_table(doc, TABLES["external_ai"])

    doc.add_page_break()
    doc.add_heading("附錄：複現指令", level=1)
    doc.add_paragraph("python scripts/build_v39d_timeline.py --root <ESG競賽根目錄>")
    doc.add_paragraph("python scripts/reproduce_v43.py --root <ESG競賽根目錄>")
    doc.save(DOCX_OUT)


def para(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), style)


def pdf_figure(path: Path, caption: str, style: ParagraphStyle, width_cm: float = 17.0) -> list:
    if not path.exists():
        return []
    with Image.open(path) as img:
        w_px, h_px = img.size
    width = width_cm * cm
    height = width * h_px / w_px
    return [
        para(caption, style),
        RLImage(str(path), width=width, height=height),
        Spacer(1, 0.25 * cm),
    ]


def build_pdf() -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "ChineseNormal",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=10,
        leading=14,
        wordWrap="CJK",
        spaceAfter=6,
    )
    h1 = ParagraphStyle(
        "ChineseH1",
        parent=styles["Heading1"],
        fontName="STSong-Light",
        fontSize=15,
        leading=19,
        spaceBefore=10,
        spaceAfter=8,
    )
    title = ParagraphStyle(
        "ChineseTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=17,
        leading=22,
        alignment=1,
        spaceAfter=14,
    )
    small = ParagraphStyle(
        "ChineseSmall",
        parent=normal,
        fontSize=8,
        leading=11,
    )

    story = [Paragraph(TITLE, title), Paragraph("隊伍與成績資訊", h1), make_pdf_table(TABLES["team"], small)]
    for heading, paragraphs in SECTIONS:
        story.append(Paragraph(heading, h1))
        for paragraph in paragraphs:
            story.append(para(paragraph, normal))
        if heading.startswith("二、"):
            story.append(Paragraph("GitHub 架構文章對應", h1))
            story.append(make_pdf_table(TABLES["github_architecture_map"], small))
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph("完整演算法層架構圖", h1))
            story.extend(pdf_figure(FIGURES["algo"], "圖 0：依 GitHub 公開架構整理的完整演算法層，從資料層一路到 deployment gate。", small))
            story.append(Paragraph("演算法層對照表", h1))
            story.append(make_pdf_table(TABLES["algorithm_layers"], small))
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph("最終欄位來源與演算法理由", h1))
            story.append(make_pdf_table(TABLES["field_algorithm"], small))
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph("最終組裝流程圖", h1))
            story.extend(pdf_figure(FIGURES["overview"], "圖 1：完整模型收斂到 v25 base，再由 v39d timeline-only 分支形成最終 v43。", small))
            story.extend(pdf_figure(FIGURES["field"], "圖 2：field-best assembly：四欄各自選擇通過 gate 的最佳演算法來源。", small))
            story.append(Paragraph("v43 vs v25 欄位異動", h1))
            story.append(make_pdf_table(TABLES["v43_vs_v25"], small))
            story.append(Spacer(1, 0.2 * cm))
        if heading.startswith("三、"):
            story.append(Paragraph("驗證與部署 gate", h1))
            story.extend(pdf_figure(FIGURES["gate"], "圖 3：所有新分支先做單欄 ablation 與 gate，未通過者不併入 private final。", small))
            story.append(Paragraph("未納入 v43 的支線", h1))
            story.append(make_pdf_table(TABLES["rejected"], small))
            story.append(Spacer(1, 0.2 * cm))
        if heading.startswith("六、"):
            story.append(Paragraph("本地驗證與 private 結果對齊", h1))
            story.append(make_pdf_table(TABLES["local_private_alignment"], small))
            story.append(Spacer(1, 0.2 * cm))
        if heading.startswith("七、"):
            story.append(Paragraph("主要腳本與角色", h1))
            story.append(make_pdf_table(TABLES["script_map"], small))
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph("必要 artifact checksum", h1))
            story.append(make_pdf_table(TABLES["checksums"], small))
            story.append(Spacer(1, 0.2 * cm))
        if heading.startswith("八、"):
            story.append(Paragraph("生成式 AI / 外部模型使用摘要表", h1))
            story.append(make_pdf_table(TABLES["external_ai"], small))
            story.append(Spacer(1, 0.2 * cm))

    story.append(PageBreak())
    story.append(Paragraph("附錄：複現指令", h1))
    story.append(para("python scripts/build_v39d_timeline.py --root <ESG競賽根目錄>", normal))
    story.append(para("python scripts/reproduce_v43.py --root <ESG競賽根目錄>", normal))

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        rightMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )
    doc.build(story)


def make_pdf_table(rows: list[list[str]], style: ParagraphStyle) -> Table:
    wrapped = [[para(str(cell), style) for cell in row] for row in rows]
    col_count = len(rows[0])
    if col_count == 2:
        widths = [5.0 * cm, 12.0 * cm]
    elif col_count == 3:
        widths = [4.2 * cm, 6.2 * cm, 6.6 * cm]
    else:
        widths = None
    table = Table(wrapped, colWidths=widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def main() -> None:
    make_figures()
    build_docx()
    build_pdf()
    print(DOCX_OUT)
    print(PDF_OUT)


if __name__ == "__main__":
    main()
